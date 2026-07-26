from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


INK = "243238"
BLUE = "D8E4F2"
BLUE_DARK = "54789A"
GREEN = "E3F0DD"
GREEN_DARK = "2F6B45"
GOLD = "F7E7B4"
CORAL = "F4D1C8"
GRAY = "ECEFF1"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(*(int(value[index : index + 2], 16) for index in (0, 2, 4)))


def set_run_font(run, *, name: str = "宋体", size: float = 10.5, bold: bool | None = None, color: str = INK) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = rgb(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, total_twips: int = 9360, widths: list[int] | None = None) -> None:
    columns = len(table.columns)
    if widths is None:
        widths = [total_twips // columns] * columns
        widths[-1] += total_twips - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_document(doc: Document) -> None:
    definitions = [
        ("Normal", 10.5, False, INK),
        ("Heading 1", 19, True, INK),
        ("Heading 2", 14, True, INK),
        ("Heading 3", 11.5, True, INK),
        ("List Bullet", 10.5, False, INK),
        ("List Number", 10.5, False, INK),
        ("Template Title", 26, True, INK),
        ("Template Note", 9.5, False, BLUE_DARK),
    ]
    for style_name, size, bold, color in definitions:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name == "Heading 1" else 4)


def add_field_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "打开 Word 后更新整个目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def append_simple_field(run, instruction_text: str, fallback: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {instruction_text} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    first = paragraph.add_run()
    set_run_font(first, size=9, color="666666")
    append_simple_field(first, "PAGE", "1")
    middle = paragraph.add_run(" / ")
    set_run_font(middle, size=9, color="666666")
    last = paragraph.add_run()
    set_run_font(last, size=9, color="666666")
    append_simple_field(last, "NUMPAGES", "1")


def add_table(doc: Document, rows: list[list[str]], widths: list[int] | None = None, font_size: float = 8.8):
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    set_table_geometry(table, widths=widths)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for column_index, cell in enumerate(table.rows[row_index].cells):
            value = values[column_index] if column_index < len(values) else ""
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            set_run_font(run, size=font_size, bold=row_index == 0)
    doc.add_paragraph()
    return table


def add_section_table(doc: Document, title: str, rows: list[list[str]], widths: list[int]) -> None:
    title_table = doc.add_table(rows=1, cols=1)
    title_table.style = "Table Grid"
    set_table_geometry(title_table, widths=[9360])
    set_repeat_table_header(title_table.rows[0])
    set_cell_shading(title_table.cell(0, 0), BLUE)
    paragraph = title_table.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_run_font(run, name="微软雅黑", size=11.5, bold=True)
    add_table(doc, rows, widths=widths, font_size=8.7)


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.replace("**", "").replace("__", "").replace("`", "").strip()


def add_text_paragraph(doc: Document, text: str, *, style: str = "Normal", color: str = INK, size: float = 10.5) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=size, color=color)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    widths = [9360 // columns] * columns
    widths[-1] += 9360 - sum(widths)
    if columns >= 3:
        widths[0] = int(9360 * 0.18)
        remaining = 9360 - widths[0]
        widths[1:] = [remaining // (columns - 1)] * (columns - 1)
        widths[-1] += 9360 - sum(widths)
    add_table(doc, [[clean_inline(cell) for cell in row] for row in rows], widths=widths, font_size=7.8 if columns >= 6 else 8.5)


def parse_markdown(doc: Document, markdown: str, *, include_title: bool = False) -> None:
    lines = markdown.splitlines()
    index = 0
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_text_paragraph(doc, " ".join(item.strip() for item in paragraph_buffer if item.strip()))
            paragraph_buffer = []

    while index < len(lines):
        raw = lines[index].rstrip()
        line = raw.strip()
        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                if code_lines:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run("\n".join(code_lines))
                    set_run_font(run, name="Consolas", size=8.3, color=GREEN_DARK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not line or line == "---":
            flush_paragraph()
            index += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows: list[list[str]] = []
            for table_line in table_lines:
                parts = [part.strip() for part in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts):
                    continue
                rows.append(parts)
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            marks, title = heading.groups()
            if len(marks) == 1 and not include_title:
                index += 1
                continue
            level = min(3, max(1, len(marks) - (0 if include_title else 1)))
            doc.add_paragraph(clean_inline(title), style=f"Heading {level}")
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            add_text_paragraph(doc, line, style="Normal")
            index += 1
            continue
        if line.startswith("- "):
            flush_paragraph()
            add_text_paragraph(doc, line[2:], style="List Bullet")
            index += 1
            continue
        if line.startswith(">"):
            flush_paragraph()
            add_text_paragraph(doc, line.lstrip(">").strip(), style="Template Note", color=BLUE_DARK, size=9.8)
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()


def add_baseline_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Cm(8.6))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", image_path.stem)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=8.8, color="666666")


def add_review_board_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Cm(16.6))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", image_path.stem)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=8.8, color="666666")


def add_status_panel(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, widths=[9360])
    set_repeat_table_header(table.rows[0])
    set_cell_shading(table.cell(0, 0), CORAL)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run(
        "当前最高 Gate：PENPOT_EDITABLE_SOURCE_READBACK_VERIFIED。"
        "Penpot 云端导出归档和用户详细设计评审仍为 PENDING；runtime_authority=false；"
        "本文、PDF、静态截图和测试均不证明运行时完成。"
    )
    set_run_font(run, size=8.8, bold=True)
    doc.add_paragraph()


def ensure_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_document(
    template: Path,
    feature_md: Path,
    ui_md: Path,
    visual_md: Path,
    decision_md: Path,
    penpot_decision_md: Path,
    baseline_main: Path,
    baseline_orders: Path,
    penpot_screen_board: Path,
    penpot_flow_board: Path,
    output: Path,
) -> None:
    doc = Document(str(template))
    clear_document_body(doc)
    style_document(doc)
    add_footer(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    title = doc.add_paragraph(style="Template Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("F-004 动物居民小镇：空间秩序与自主作业基础")
    set_run_font(title_run, name="微软雅黑", size=20, bold=True)
    subtitle = doc.add_paragraph(style="Template Note")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("产品与设计重基线 V1.1 · Penpot 可编辑移交 · 720×1280 竖屏 · runtime_authority=false")
    set_run_font(subtitle_run, name="微软雅黑", size=9, color=BLUE_DARK)
    add_status_panel(doc)

    add_section_table(
        doc,
        "版本控制",
        [
            ["文档编号", "COA-F004-RESIDENT.1", "当前版本", "V1.1 REVIEW"],
            ["功能 ID", "F-004 / F004-RESIDENT.1", "文档状态", "DESIGN_REBASELINE_REVIEW"],
            ["制作人 / 主策划", "Codex /root", "运行时权限", "false"],
            ["创建日期", "2026-07-26", "目标平台", "Mobile / 720×1280 portrait"],
            ["默认语言", "zh-CN；设置可切 en", "配置版本", "仅建议，未授权建表"],
            ["Penpot 源", "output/penpot/F004-RESIDENT.1/README.md", "Penpot Gate", "READBACK VERIFIED / EXPORT + USER REVIEW PENDING"],
            ["正式决策", "PD-002 rebaseline + PD-003 Penpot source", "评审结论", "待用户审阅"],
        ],
        [1600, 3280, 1600, 2880],
    )
    add_section_table(
        doc,
        "版本历史",
        [
            ["版本", "编写人", "审核人", "批准人", "日期", "更新内容"],
            ["V1.0", "Codex /root", "待用户评审", "Pending", "2026-07-26", "完成 A-H 正式设计覆盖、RAG/控制面续接与 Figma 阻塞证据收口。"],
            ["V1.1", "Codex /root", "待用户评审", "Pending", "2026-07-26", "正式改用 Penpot；认证云端文件已导入 8 张界面与 4 张流程，登记 14 个对象引用并完成代表性嵌套矢量回读。"],
        ],
        [900, 1350, 1350, 1150, 1350, 3260],
    )
    add_section_table(
        doc,
        "开发计划",
        [
            ["责任方", "负责人", "当前状态", "下一 Gate", "完成条件"],
            ["产品/设计", "Codex /root", "重基线已形成，待审阅", "DESIGN_REBASELINE_APPROVED", "占地、居民、车辆与减法规则批准"],
            ["UI/UX/主美", "Codex /root", "Penpot 云端对象回读已验证；用户评审待办", "VISUAL_CONTRACT_APPROVED", "云端导出归档、全状态与评审面批准"],
            ["程序", "未授权", "未开始", "RUNTIME_SLICE_APPROVED", "一个真实 720×1280 代表切片及行为证据"],
            ["扩面", "未授权", "禁止开始", "SCALE_OUT_APPROVED", "代表切片无 BLOCKER/MATERIAL"],
        ],
        [1350, 1800, 1950, 1950, 2310],
    )

    doc.add_page_break()
    doc.add_paragraph("目录", style="Heading 1")
    toc = doc.add_paragraph()
    add_field_toc(toc)
    note = doc.add_paragraph(style="Template Note")
    run = note.add_run("目录在 PDF 导出前由 Word 更新，覆盖 1–3 级标题。")
    set_run_font(run, size=9.5, color=BLUE_DARK)
    doc.add_page_break()

    doc.add_paragraph("0. 制作人只读基线", style="Heading 1")
    add_text_paragraph(
        doc,
        "F003-FARM.2 仅保留为已接受可玩原型基线；旧 F004-DISTRICT.1 为 22% 的历史 FORMAL DESIGN HOLD / FIGMA，"
        "其频繁点击、固定站点和分区面板方向与本次正式决策冲突，进入 revise/superseded 审计。"
    )
    add_text_paragraph(
        doc,
        "以下截图来自真实 720×1280 运行证据，仅用于确认旧基线问题；不是新方向 UI、视觉来源或运行时验收。"
    )
    add_baseline_figure(doc, baseline_main, "图 0-1：旧 F003 主地图真实运行基线。问题包括占地不统一、道路断裂、居民不是任务执行者。")
    add_baseline_figure(doc, baseline_orders, "图 0-2：旧 F003 订单面板真实运行基线。订单仍是卡片/按钮，而不是世界车辆事件。")

    doc.add_paragraph("0.1 Penpot 可编辑设计移交", style="Heading 1")
    add_text_paragraph(
        doc,
        "下列画板由项目内 SVG 可编辑源渲染，并已导入认证 Penpot 文件 "
        "CityOfAnimals / F004-RESIDENT.1（file bd31e32d-d69f-81e2-8008-62cc67c1eeda；"
        "page bd31e32d-d69f-81e2-8008-62cc67c1eedb）。"
        "文件已重新打开，8 个界面组、4 个流程组、14 个根/组对象引用及代表性嵌套矢量完成回读。"
        "本地 SVG 是导入源与备份；PNG/PDF 仅为派生评审预览，云端导出归档与用户详细评审仍待完成。"
    )
    add_review_board_figure(
        doc,
        penpot_screen_board,
        "图 0-3：Penpot 界面导入源总览。覆盖 8 个命名界面组，包括 720×1280 主地图、建造、居民、车辆订单、系统状态与组件 Token。",
    )
    add_review_board_figure(
        doc,
        penpot_flow_board,
        "图 0-4：Penpot 流程导入源总览。覆盖代表性 UE、居民状态机、车辆订单状态机与空间放置合法性。",
    )

    parse_markdown(doc, feature_md.read_text(encoding="utf-8"))
    appendix_a = doc.add_paragraph("附录 A：制作人决策原文与处置", style="Heading 1")
    appendix_a.paragraph_format.page_break_before = True
    parse_markdown(doc, decision_md.read_text(encoding="utf-8"))
    appendix_b = doc.add_paragraph("附录 B：主页面 UI/UX 优先级", style="Heading 1")
    appendix_b.paragraph_format.page_break_before = True
    parse_markdown(doc, ui_md.read_text(encoding="utf-8"))
    appendix_c = doc.add_paragraph("附录 C：视觉质量合同", style="Heading 1")
    appendix_c.paragraph_format.page_break_before = True
    parse_markdown(doc, visual_md.read_text(encoding="utf-8"))
    appendix_d = doc.add_paragraph("附录 D：Penpot 正式设计源决策", style="Heading 1")
    appendix_d.paragraph_format.page_break_before = True
    parse_markdown(doc, penpot_decision_md.read_text(encoding="utf-8"))

    ensure_update_fields(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"built={output}")
    print(f"sha256={hashlib.sha256(output.read_bytes()).hexdigest()}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--feature", type=Path, required=True)
    parser.add_argument("--ui", type=Path, required=True)
    parser.add_argument("--visual", type=Path, required=True)
    parser.add_argument("--decision", type=Path, required=True)
    parser.add_argument("--penpot-decision", type=Path, required=True)
    parser.add_argument("--baseline-main", type=Path, required=True)
    parser.add_argument("--baseline-orders", type=Path, required=True)
    parser.add_argument("--penpot-screen-board", type=Path, required=True)
    parser.add_argument("--penpot-flow-board", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_document(
        args.template,
        args.feature,
        args.ui,
        args.visual,
        args.decision,
        args.penpot_decision,
        args.baseline_main,
        args.baseline_orders,
        args.penpot_screen_board,
        args.penpot_flow_board,
        args.output,
    )


if __name__ == "__main__":
    main()
