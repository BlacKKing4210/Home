from docx import Document

REFERENCE = r"C:\Users\76398\.codex\skills\artifact-template-game-feature-design-general\assets\reference.docx"

document = Document(REFERENCE)
print(f"PARAGRAPHS {len(document.paragraphs)} TABLES {len(document.tables)}")
for index, paragraph in enumerate(document.paragraphs):
    if paragraph.text.strip() and index >= 93:
        print(f"P{index}: {paragraph.style.name}: {paragraph.text[:160]}")

for table_index, table in enumerate(document.tables):
    print(f"T{table_index} {len(table.rows)}x{len(table.columns)}")
    for row_index, row in enumerate(table.rows[:8]):
        cells = [cell.text.replace("\n", " / ")[:120] for cell in row.cells]
        print(f"  R{row_index}: {' || '.join(cells)}")
