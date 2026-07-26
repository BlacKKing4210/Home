#!/usr/bin/env python3
"""Build the formal F003-FARM.2 DOCX from the retained general template."""

from __future__ import annotations

import shutil
from io import BytesIO
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[3]
REFERENCE = Path(
    r"C:\Users\76398\.codex\skills\artifact-template-game-feature-design-general"
    r"\assets\reference.docx"
)
OUTPUT = PROJECT_ROOT / "output" / "documents" / "F003-FARM.2" / "CityOfAnimals_F003_Farm_Town_Foundation_V2.docx"
FIGMA_URL = "https://www.figma.com/design/uU2Oek5RqFb19CPoGl48lC/Untitled"
FIGMA_PREVIEW = PROJECT_ROOT / "output" / "figma" / "F003-FARM.2" / "city_of_animals_f003_farm2_ue-preview.png"
ART_MANIFEST = PROJECT_ROOT / "output" / "art" / "ART-003-FARM2" / "split-manifest.json"

USABLE_WIDTH_DXA = 10200
HEADER_FILL = "BDD7EE"
SECTION_FILL = "D9E2F3"
GREEN = RGBColor(0x00, 0xA6, 0x51)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xE5, 0x00, 0x00)
GRAY = RGBColor(0x7F, 0x7F, 0x7F)


def clear_body(document: Document) -> None:
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_geometry(table, ratios: list[float]) -> None:
    widths = [round(USABLE_WIDTH_DXA * ratio / sum(ratios)) for ratio in ratios]
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_runs(paragraph, *, bold: bool | None = None, color: RGBColor | None = None, size: float | None = None) -> None:
    for run in paragraph.runs:
        if bold is not None:
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)


def add_table(document: Document, headers: list[str], rows: list[list[str]], ratios: list[float]) -> object:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, HEADER_FILL)
        style_runs(cell.paragraphs[0], bold=True)
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        set_cant_split(row)
        for index, value in enumerate(values):
            row.cells[index].text = str(value)
    set_table_geometry(table, ratios)
    document.add_paragraph()
    return table


def add_section_band(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = text
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    shade_cell(table.cell(0, 0), SECTION_FILL)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_runs(paragraph, bold=True, size=12)
    set_table_geometry(table, [1])


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_colored_note(document: Document, label: str, text: str, color: RGBColor) -> None:
    paragraph = document.add_paragraph()
    lead = paragraph.add_run(f"{label}：")
    lead.bold = True
    lead.font.color.rgb = color
    paragraph.add_run(text)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0070C0")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "目录将在 Word 字段刷新后生成。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run1 = OxmlElement("w:r")
    run1.append(begin)
    run2 = OxmlElement("w:r")
    run2.append(instruction)
    run3 = OxmlElement("w:r")
    run3.append(separate)
    run4 = OxmlElement("w:r")
    run4.append(display)
    run5 = OxmlElement("w:r")
    run5.append(end)
    for run in (run1, run2, run3, run4, run5):
        paragraph._p.append(run)


def crop_stream(box: tuple[int, int, int, int]) -> BytesIO:
    with Image.open(FIGMA_PREVIEW) as image:
        crop = image.crop(box)
        stream = BytesIO()
        crop.save(stream, format="PNG", optimize=True)
        stream.seek(0)
        return stream


def add_figure(document: Document, box: tuple[int, int, int, int], width_inches: float, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    inline_shape = run.add_picture(crop_stream(box), width=Inches(width_inches))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", f"{caption}。可编辑源见 {FIGMA_URL}")
    caption_paragraph = document.add_paragraph(caption, style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_fidelity_register(document: Document, figure_type: str, nodes: str, coverage: str) -> None:
    table = add_table(
        document,
        ["图类型", "Figma/FigJam 可编辑链接", "版本 / 日期", "负责人"],
        [[figure_type, FIGMA_URL, "V1.0 / 2026-07-24", "Codex /root"]],
        [1.1, 3.0, 1.2, 1.0],
    )
    paragraph = document.add_paragraph()
    paragraph.add_run("节点：").bold = True
    paragraph.add_run(nodes)
    paragraph.add_run("；覆盖：").bold = True
    paragraph.add_run(coverage)


def build_document() -> None:
    if not REFERENCE.exists() or not FIGMA_PREVIEW.exists() or not ART_MANIFEST.exists():
        raise FileNotFoundError("Required template/Figma/ART evidence is missing")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    document = Document(OUTPUT)
    clear_body(document)
    properties = document.core_properties
    properties.title = "动物小镇生产经营地基 V2"
    properties.subject = "CityOfAnimals F003-FARM.2 通用复杂功能策划案"
    properties.author = "Codex /root"
    properties.last_modified_by = "Codex /root"
    properties.comments = "Commercial-safe original farm-town implementation plan."

    title = document.add_paragraph("动物小镇生产经营地基 V2", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = document.add_paragraph("COA-F003-FARM.2 · 通用复杂功能策划案 · V1.0", style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(10)
    document.add_paragraph(
        "本文定义一套面向 720 x 1280 竖屏手游的动物小镇经营地基：以场景对象直接传达生产、仓储、订单与扩张，"
        "在商业安全的原创实现边界内学习成熟农场经营品类的系统结构，不复制其资产、地图、UI、数值或文案。"
    )

    add_section_band(document, "版本控制")
    add_table(
        document,
        ["文档编号", "COA-F003-FARM.2", "当前版本", "V1.0"],
        [
            ["功能 ID", "F-003", "文档状态", "设计评审通过 / 工程门禁待开启"],
            ["主策划", "Codex /root", "制作策划", "Codex /root"],
            ["制作程序", "Codex /root（待工程回执）", "制作美术", "Codex /root"],
            ["创建日期", "2026-07-24", "计划完成", "按工程迭代回执"],
            ["Figma 源", FIGMA_URL, "配置版本", "F003-FARM.2 planned CSV contract"],
            ["备注", "Figma 与 ART-003 门禁已通过", "评审结论", "进入 DOCX/PDF 评审"],
        ],
        [1.0, 2.6, 1.0, 2.6],
    )
    add_section_band(document, "版本历史")
    add_table(
        document,
        ["版本", "编写人", "审核人", "批准人", "日期", "更新内容"],
        [["V1.0", "Codex /root", "Codex /root", "Codex /root", "2026-07-24", "建立经营地基、Figma、ART-003 与工程验收合同"]],
        [0.8, 1.0, 1.0, 1.0, 1.1, 2.3],
    )
    add_section_band(document, "开发计划")
    add_table(
        document,
        ["责任方", "负责人", "开始日期", "预计天数", "完成条件"],
        [
            ["策划", "Codex /root", "2026-07-24", "已完成", "规则、配置合同、Figma 与文档评审通过"],
            ["美术/UI", "Codex /root", "2026-07-24", "已完成准备", "26 项 ART-003 候选可追溯且 NOT_RUNTIME"],
            ["程序", "Codex /root", "工程回执后", "三轮迭代", "模型、地图、资源、UI 与行为证据通过"],
        ],
        [1.0, 1.2, 1.2, 1.0, 3.0],
    )
    toc_title = document.add_paragraph("目录", style="Title")
    toc_title.paragraph_format.page_break_before = True
    add_toc(document)
    document.add_paragraph("打开 Word 后可选择“更新整个目录”；本交付在导出 PDF 前由 Word 刷新字段。").runs[0].font.color.rgb = GRAY

    document.add_heading("1. 术语缩写与修订标记", level=1)
    add_table(
        document,
        ["术语 / 标记", "含义"],
        [
            ["F003-FARM.2", "当前有效的动物小镇生产经营地基版本；旧 F003-FARM.1 仅保留历史。"],
            ["粮仓 / granary", "只存作物与饲料；容量独立。"],
            ["货物屋 / storehouse", "只存动物产品、机器成品和材料；容量独立。"],
            ["成品位", "机器完成后的可收取槽；未收取时阻塞后续完成。"],
            ["邻里委托", "本地、动态、多物品请求；完成后原子扣货和发奖。"],
            ["橡果集市", "本地 NPC 余货出售渠道；价值低于目标委托，用于释放仓位。"],
            ["ART-003", "26 项项目自有资源的来源、透明化、尺寸、原点和联系表门禁。"],
            ["NOT_RUNTIME", "候选已准备但不得被 Godot 直接加载；须由独立工程回执晋级。"],
        ],
        [1.8, 5.4],
    )
    add_colored_note(document, "红色", "强约束与不可违反边界。", RED)
    add_colored_note(document, "蓝色", "F003-FARM.2 新增或替代内容。", BLUE)
    add_colored_note(document, "绿色", "准确配置表、字段、公式与数据源。", GREEN)
    add_colored_note(document, "灰色", "解释性说明，不形成额外批准。", GRAY)

    document.add_heading("2. 设计目的", level=1)
    document.add_heading("2.1 主要目标", level=2)
    add_bullets(
        document,
        [
            "让玩家在十几秒内通过田地、动物、机器、仓储、委托和市场完成一次看得见的经营取舍。",
            "把固定卡片式原型改造成大于一屏、可拖拽、可辨识 18 个以上场景对象的 45°动物小镇。",
            "以真实库存种子、双仓储、队列阻塞和原子事务建立可持续扩展的经济地基。",
            "用图标、建筑状态和对象反馈优先传达，文字只解释选择、数量和异常。",
        ],
    )
    document.add_heading("2.2 次要目标", level=2)
    add_bullets(
        document,
        [
            "默认简体中文，设置中可切换英文并持久化。",
            "为后续火车多箱货运、空运大订单、访客服务、区域扩张和协作预留接口。",
            "保证调参来源位于 `config/tables/*.csv`，不把数值埋在视图逻辑。",
        ],
    )
    document.add_heading("2.3 非目标", level=2)
    add_bullets(
        document,
        [
            "不复制 Hay Day、Township 或其他商业游戏的角色、建筑、图标、地图、UI 布局、文案、配方、数值和商业资产。",
            "本版本不实现联网玩家市场、联盟、社交、付费、广告、火车或空运；这些必须单独立项。",
            "本策划案和静态评审图不证明 Godot 运行时完成。",
        ],
    )

    document.add_heading("3. 功能概述", level=1)
    document.add_paragraph(
        "核心循环：观察场景状态 -> 连续收获并获得净增殖 -> 消耗种子再次播种 -> 把作物送往饲料、动物和机器 -> "
        "收取受仓储约束的产物 -> 在邻里委托与余货集市之间选择 -> 获得金币/声望 -> 建设、扩容并解锁新的生产网络。"
    )
    add_table(
        document,
        ["时间层", "玩家主要决策", "可感知反馈"],
        [
            ["5-30 秒", "收获、播种、选作物、查看容量", "田格连续变化、库存净增长、成熟/阻塞图标"],
            ["1-5 分钟", "饲料、动物、机器队列排程", "计时、成品位、队列空位与来源导航"],
            ["3-15 分钟", "三张动态委托、余货出售、建设选择", "多物品缺口、金币/声望、可建地块"],
            ["小时级", "扩容、解锁建筑和地图扩张", "新的生产链、空间和订单组合"],
        ],
        [1.2, 3.0, 3.0],
    )

    document.add_heading("4. 系统框架", level=1)
    add_fidelity_register(
        document,
        "系统框架图",
        "`city_of_animals_f003_farm2_ue 1 / 00_System_Framework`",
        "资源流、时间层、阻塞点、订单、市场、建设与后续接口",
    )
    add_figure(document, (40, 40, 1440, 940), 6.6, "图 1 生产经营系统框架（评审导出；可编辑源见 Figma）")
    add_table(
        document,
        ["层", "职责", "不可越界"],
        [
            ["玩家/场景层", "选择田地、动物、机器、委托、市场和建设对象；发送意图。", "不得直接改库存、计时和奖励。"],
            ["权威模型层", "预检、原子扣发、状态机、队列、容量、生成种子和保存脏标记。", "不得依赖可见控件作为事实来源。"],
            ["配置/资源层", "CSV、locale、ART-003 正式晋级资源和世界坐标。", "不得在视图脚本硬编码调参和玩家文本。"],
        ],
        [1.2, 3.2, 2.8],
    )

    document.add_heading("5. UE 流程图", level=1)
    add_fidelity_register(
        document,
        "玩家 UE 流程图",
        "`city_of_animals_f003_farm2_ue 1 / 01_Player_UE_Flow`",
        "进入、收获、播种、养殖、加工、订单、扩容以及阻塞/中断返回",
    )
    add_figure(document, (1500, 40, 3340, 940), 6.8, "图 2 一次完整可感知经营循环（评审导出；可编辑源见 Figma）")
    add_table(
        document,
        ["起点", "玩家动作 / 系统事件", "条件", "终点", "失败 / 返回去向"],
        [
            ["主地图", "划过成熟田格收获", "粮仓可容纳整格产量", "空田 + 作物盘入口", "粮仓满：保持成熟并打开容量来源"],
            ["作物盘", "选作物并划过空田", "库存足够且作物已解锁", "生长田", "缺种：停在最后成功田格并提示来源"],
            ["混粮坊", "把饲料配方加入队列", "队列有空位且原料满足", "加工中", "队列满/缺料：不扣除并标记缺口"],
            ["动物栏", "喂养", "有对应饲料且动物空腹", "生产中", "饲料不足：导航到混粮坊/来源"],
            ["动物栏", "收取鸡蛋或牛奶", "货物屋可容纳完整产量", "空腹", "货物屋满：保持可收取"],
            ["机器", "加入配方", "输入足够且队列未满", "排队/加工", "不满足时原子失败"],
            ["机器", "收取成品", "货物屋有空间", "下一队列项启动", "成品位继续阻塞"],
            ["委托板", "提交一张请求", "所有需求一次性满足", "金币/声望 + 新请求", "任一不足：不扣货、不换单"],
            ["橡果集市", "选择物品和数量并确认", "库存满足", "金币 + 仓位释放", "取消或不足：保持原库存"],
            ["设置", "切换语言", "选择 zh-CN 或 en", "立即刷新并保存", "保存失败：保留旧选择并提示"],
        ],
        [1.2, 2.0, 1.8, 1.5, 2.1],
    )

    document.add_heading("6. 参考视频", level=1)
    paragraph = document.add_paragraph("制作人补充视频：")
    add_hyperlink(paragraph, "YouTube - zgAysNgvDuk", "https://www.youtube.com/watch?v=zgAysNgvDuk")
    add_bullets(
        document,
        [
            "只参考高频作物循环如何同时牵动库存、仓储、生产和出售吞吐；播放器内容未逐帧核验。",
            "系统事实优先使用 Supercell 官方 Hay Day 页面与支持资料；完整来源和证据等级见 `docs/research/2026-07-hay-day-moment-to-moment-gap-study.md`。",
            "所有落地规则均重新命名、重新取值、重新组织内容和视觉表达；视频/截图不是资产来源。",
        ],
    )

    document.add_heading("7. 配置表调整", level=1)
    add_colored_note(
        document,
        "配置归属",
        "所有可调项进入下列 `config/tables/*.csv`；运行时生成器、状态机和视图只能读取，不得复制常量。",
        GREEN,
    )
    config_rows = [
        ["items", "f003_v2_items.csv::item_id/category/storage_type", "id/enum", "库存、保存、配方、委托和 locale 关联；首轮 9 个原创物品"],
        ["items", "f003_v2_items.csv::initial_amount", "int", "初始库存；通过前 10 分钟无硬锁模拟"],
        ["items", "f003_v2_items.csv::market_coin_value", "int", "NPC 市场价值；与委托奖励分离"],
        ["crops", "f003_v2_crops.csv::grow_seconds", "float/s", "8 / 25 / 60；原创 MVP 时长"],
        ["crops", "f003_v2_crops.csv::plant_cost", "int", "首轮均为 1"],
        ["crops", "f003_v2_crops.csv::harvest_yield", "int", "3 / 3 / 4，必须大于 plant_cost"],
        ["storage", "f003_v2_storage.csv::base_capacity", "int", "粮仓 48；货物屋 40"],
        ["storage", "f003_v2_storage.csv::capacity_per_level", "int", "粮仓 12；货物屋 10"],
        ["storage", "f003_v2_storage.csv::upgrade_costs", "id-count", "引用 items 中的原创建材"],
        ["recipes", "f003_v2_recipes.csv::input_items/output_item_id", "id-count/id", "入队时原子扣除；产出进入成品位"],
        ["recipes", "f003_v2_recipes.csv::duration_seconds", "float/s", "12-50；配方独立"],
        ["recipes", "f003_v2_recipes.csv::queue_slots", "int", "由 buildings.csv 的机器配置提供 2-3 格"],
        ["animals", "f003_v2_animals.csv::animal_count", "int", "鸡 3；牛 2"],
        ["animals", "f003_v2_animals.csv::feed_cost", "int", "鸡 1；牛 2"],
        ["animals", "f003_v2_animals.csv::output_count", "int", "鸡蛋 2；牛奶 2"],
        ["animals", "f003_v2_animals.csv::duration_seconds", "float/s", "鸡 18；牛 30"],
        ["buildings", "f003_v2_buildings.csv::world_x/world_y", "float", "大地图世界坐标，不写进视图"],
        ["buildings", "f003_v2_buildings.csv::footprint_w/footprint_h", "int", "等距地块占用"],
        ["buildings", "f003_v2_buildings.csv::asset_id", "id", "ART-003 正式晋级后的运行时 ID"],
        ["requests", "f003_v2_requests.csv::slot_count", "int", "首轮 3"],
        ["requests", "f003_v2_requests.csv::min_item_types/max_item_types", "int", "1 / 2"],
        ["requests", "f003_v2_requests.csv::discard_refresh_seconds", "float/s", "30；原创 MVP 值"],
        ["requests", "f003_v2_requests.csv::seed_salt", "string", "保存并可复现，重进不换单"],
        ["world", "f003_v2_world.csv::world_width/world_height", "int/px", "1800 / 1700"],
        ["world", "f003_v2_world.csv::camera_pan_threshold_px", "float/px", "12"],
        ["world", "f003_v2_world.csv::field_drag_step_px", "float/px", "24"],
        ["world", "f003_v2_world.csv::starter_field_count", "int", "12"],
        ["world", "f003_v2_world.csv::autosave_interval_seconds", "float/s", "10"],
        ["locale", "f003_v2_locale.csv::locale_key/zh_CN/en", "string", "所有玩家文本；首次 zh-CN，可选 en"],
        ["preferences", "user://city_of_animals_preferences.cfg::language_code", "string", "语言选择持久化"],
    ]
    config_table = add_table(document, ["域", "准确字段", "类型", "规则 / 默认"], config_rows, [1.0, 3.0, 1.1, 2.5])
    for row in config_table.rows[1:]:
        style_runs(row.cells[1].paragraphs[0], color=GREEN)
    document.add_paragraph(
        "容量公式：capacity = base_capacity + (upgrade_level - 1) * capacity_per_level。"
    )
    document.add_paragraph(
        "首轮配方：golden_sprig*2 + cloud_bean*1 -> leafy_feed（12s）；golden_sprig*3 + spotted_egg*1 -> hearth_loaf（35s）；"
        "cloud_milk*2 -> soft_cream（40s）；root_carrot*3 + cloud_bean*1 -> root_preserve（50s，后续同批候选）。"
    )

    document.add_heading("8. 系统逻辑", level=1)
    document.add_heading("8.1 系统整体逻辑", level=2)
    add_fidelity_register(
        document,
        "系统整体逻辑 / UE 交付",
        "`00_System_Framework` + `01_Player_UE_Flow`",
        "玩家操作、场景反馈、模型事务、配置与保存边界",
    )
    add_figure(document, (40, 40, 3340, 940), 6.8, "图 3 系统框架与玩家 UE 双视图（评审导出）")
    add_table(
        document,
        ["泳道", "输入", "处理", "输出"],
        [
            ["玩家/交互", "点选、拖拽、连续划过、确认/取消", "只发送意图并展示模型结果", "主动作、缺口、成功/阻塞反馈"],
            ["权威模型", "意图、当前状态、Unix 时间", "预检 -> 原子提交 -> 事件 -> 保存脏标记", "库存、状态、奖励、完成时间"],
            ["配置/持久化", "CSV、locale、asset_id、存档 schema", "解析、校验、迁移、可复现生成", "调参值、语言、重进一致性"],
        ],
        [1.2, 1.8, 2.6, 2.0],
    )

    mechanisms = [
        ("8.2 作物种子经济", [
            "状态：EMPTY -> GROWING -> READY -> EMPTY。",
            "播种先检查解锁与种子库存，再逐格扣 plant_cost；已成功田格不因后续缺种回滚。",
            "收获只在粮仓能容纳完整 harvest_yield 时提交；否则保持 READY，不部分收取。",
            "首轮 12 格金穗草从 12 种子播满后为 0，收获 36，净增长 24。",
        ]),
        ("8.3 双仓储", [
            "作物/饲料进入粮仓；动物产品/机器成品/材料进入货物屋。",
            "容量检查属于模型事务前置条件；页面只显示当前相关仓储。",
            "仓满时保留成熟田、动物产物或机器成品，不溢出、不丢失。",
        ]),
        ("8.4 饲料与动物", [
            "状态：HUNGRY -> PRODUCING -> READY -> HUNGRY。",
            "喂养时一次性扣饲料并记录 ready_at；返回前台只推进状态，不自动收取。",
            "货物屋满时保持 READY；成功收取才发产物并回到 HUNGRY。",
        ]),
        ("8.5 机器队列与成品位", [
            "状态：IDLE -> PROCESSING -> OUTPUT_READY -> IDLE/PROCESSING。",
            "成功入队只扣一次输入；队列满或原料不足完全不扣。",
            "成品位未收取时，后续队列不得穿透完成；收取后立即启动下一项。",
        ]),
        ("8.6 委托与市场事务", [
            "委托一次性验证全部需求、扣除全部物品、发金币/声望、记录事件，再生成下一请求。",
            "任一步失败都不扣货、不发奖、不换单；丢弃后保存 refresh_at 与生成种子。",
            "橡果集市按选择数量原子扣货并发金币；价值低于委托，承担释放仓位的安全阀功能。",
        ]),
        ("8.7 保存、迁移与语言", [
            "存档至少保存 schema、语言、货币/声望、双库存、田地、机器、动物、委托、建筑、相机位置。",
            "旧 F003-FARM.1 迁移到安全初始状态，不把缺失旧字段解释成已扣库存。",
            "首次 zh-CN；en 可选；切换立即刷新并保存，失败则保留旧选择。",
        ]),
    ]
    for heading, bullets in mechanisms:
        document.add_heading(heading, level=2)
        add_bullets(document, bullets)

    document.add_heading("8.8 状态与边界矩阵", level=2)
    add_table(
        document,
        ["状态 / 场景", "进入条件", "允许操作", "系统处理", "页面 / 反馈", "退出条件"],
        [
            ["缺种", "下一格库存不足", "换作物/去来源", "不透支；保留已成功田格", "最后一格抖动 + 缺口图标", "有可播种作物"],
            ["粮仓满", "完整产量放不下", "开粮仓/出售", "成熟田保持", "田地成熟 + 容量脉冲", "释放足够空间"],
            ["货物屋满", "动物/机器产物放不下", "开仓/交委托/出售", "保留可收取", "产物图标 + 仓满", "释放空间"],
            ["机器队列满", "尝试追加配方", "收成品/等待", "不扣料", "禁用 CTA + 槽位边框", "有空槽"],
            ["缺原料", "配方/委托预检失败", "跳转来源/关闭", "不扣除", "图标数量缺口", "库存满足"],
            ["请求刷新", "丢弃请求", "看其他槽", "保存 refresh_at", "倒计时", "倒计时结束"],
            ["建筑锁定", "等级/金币不足", "查看需求/关闭", "不建造", "工地/锁图标", "条件满足"],
            ["中断/后台", "系统通知", "无", "保存时间戳并重算", "不弹失败", "回前台"],
            ["旧存档", "schema 低于 V2", "继续", "安全迁移并记录", "一次性提示", "新 schema 保存"],
            ["无可行动项", "全部等待且不可播种", "查看计时/市场/建设", "防硬锁救援检测", "导航到最近完成", "出现可操作对象"],
        ],
        [1.1, 1.4, 1.4, 1.6, 1.7, 1.2],
    )

    document.add_heading("9. UI 界面及子玩法", level=1)
    document.add_heading("9.1 页面清单", level=2)
    add_table(
        document,
        ["页面 ID / 名称", "入口", "退出 / 返回", "主要状态", "Figma 节点"],
        [
            ["F003-MAP 主地图", "启动/返回", "设置/系统退出", "idle/pan/selected/blocked/return", "02_Main_Map_720x1280"],
            ["F003-CROP 作物盘", "点击空田/播种工具", "选作物/点外部", "available/locked/empty", "03_Object_Interactions"],
            ["F003-CONTEXT 情境层", "选择对象", "下滑/点外部/完成", "primary/disabled/success", "03_Object_Interactions"],
            ["F003-MACHINE 机器层", "选择机器", "返回地图", "idle/processing/full/output", "03_Object_Interactions"],
            ["F003-REQUESTS 委托", "选择委托板", "返回地图", "ready/missing/refresh/success", "03_Object_Interactions"],
            ["F003-MARKET 集市", "选择市场", "取消/确认", "select/quantity/confirm/empty", "02_Main_Map_720x1280"],
            ["F003-SETTINGS 设置", "齿轮", "保存返回", "zh-CN/en/reduced-motion", "05_Settings_720x1280"],
        ],
        [1.7, 1.2, 1.3, 1.8, 2.0],
    )
    document.add_heading("9.2 主地图与对象交互", level=2)
    add_fidelity_register(
        document,
        "页面说明图",
        "`02_Main_Map_720x1280` + `03_Object_Interactions`",
        "720 x 1280 大地图、对象入口、田地/动物/机器/委托交互",
    )
    add_figure(document, (40, 1000, 760, 2250), 3.7, "图 4 720 x 1280 主地图线框（评审导出）")
    add_figure(document, (820, 1000, 3340, 2250), 6.8, "图 5 四类对象核心交互（评审导出）")
    add_bullets(
        document,
        [
            "P0：当前可操作对象状态、选中对象、一个主动作、相关物品数量。",
            "P1：金币、声望/等级、相关仓储容量、委托目标和相机方向。",
            "P2：队列槽、建设需求、最近完成对象导航、语言/设置。",
            "P3：装饰、环境叙事和次要居民气泡。",
            "明确移除：大段说明卡、同时常驻两个巨大容量条、多个同权 CTA、以文字代替建筑状态。",
        ],
    )
    document.add_heading("9.3 状态、恢复与设置", level=2)
    add_fidelity_register(
        document,
        "状态与设置图",
        "`04_States_And_Edges` + `05_Settings_720x1280`",
        "仓满、缺料、队列、市场返回、中断、语言切换和减弱动态效果",
    )
    add_figure(document, (40, 2320, 1750, 3300), 6.6, "图 6 边界状态与恢复路径（评审导出）")
    add_figure(document, (1800, 2320, 2520, 3620), 3.6, "图 7 设置与语言页面（评审导出）")
    add_table(
        document,
        ["元素 ID / 名称", "玩家含义", "显示条件", "交互行为", "数据源", "异常 / 空状态", "归属模块"],
        [
            ["HUD-COIN", "金币", "地图常驻", "只读", "town_model.coins", "加载态占位", "HUD"],
            ["HUD-RENOWN", "等级/声望", "地图常驻", "点按查看进度", "town_model.renown", "旧存档迁移提示", "HUD"],
            ["HUD-STORAGE", "当前相关仓储", "选中可产出对象", "打开仓储", "inventory + storage config", "near-full/full", "HUD"],
            ["OBJ-READY", "可收取", "对象 READY", "收取", "authoritative state", "仓满保持", "World"],
            ["CTX-PRIMARY", "当前主动作", "选择对象", "点击提交意图", "selected object + model", "禁用并显示缺口", "Context"],
            ["QUEUE-SLOTS", "机器排程", "打开机器", "添加/收取", "machine queue", "full/output blocked", "Machine"],
            ["REQ-CARDS", "三张委托", "打开委托板", "提交/丢弃/导航", "request model", "refreshing/missing", "Request"],
            ["MARKET-SELL", "余货出售", "打开集市", "选物/数量/确认", "inventory + item value", "无可售物空箱", "Market"],
            ["SET-LANGUAGE", "语言", "打开设置", "选 zh-CN/en", "preferences.language_code", "保存失败保留旧值", "Settings"],
            ["SET-MOTION", "减弱动态", "打开设置", "开关", "preferences.reduced_motion", "停止无限摇摆/强闪", "Settings"],
        ],
        [1.1, 1.1, 1.1, 1.3, 1.5, 1.4, 1.0],
    )

    document.add_heading("10. 相关需求", level=1)
    document.add_heading("10.1 美术资源需求", level=2)
    asset_ids = [
        ("plot_wheat_ready_v1", "成熟田块", "田地", "P0"),
        ("plot_clover_ready_v1", "成熟田块", "田地", "P1"),
        ("plot_sunflower_ready_v1", "成熟田块", "田地", "P1"),
        ("plot_carrot_ready_v1", "成熟田块", "田地", "P0"),
        ("orchard_apple_v1", "果园", "后续内容", "P2"),
        ("plot_tea_bush_v1", "灌木田", "后续内容", "P2"),
        ("building_feedworks_v1", "混粮坊", "机器", "P0"),
        ("pen_sheep_v1", "羊栏", "后续养殖", "P2"),
        ("building_chicken_coop_v1", "小鸡舍", "养殖", "P0"),
        ("pen_pig_v1", "猪栏", "后续养殖", "P2"),
        ("building_granary_v1", "粮仓", "仓储", "P0"),
        ("building_dairy_v1", "乳坊", "机器", "P0"),
        ("building_preserve_v1", "制酱坊", "后续机器", "P1"),
        ("building_textile_v1", "纺织坊", "后续机器", "P1"),
        ("building_juice_press_v1", "榨汁坊", "后续机器", "P1"),
        ("building_bakery_industrial_v1", "工业面包房", "机器候选", "P1"),
        ("building_chicken_coop_v2", "大鸡舍", "养殖候选", "P1"),
        ("building_storehouse_v1", "货物屋", "仓储", "P0"),
        ("building_grainworks_v1", "谷物加工坊", "机器候选", "P1"),
        ("building_bakery_shop_v1", "面包店", "机器", "P0"),
        ("building_roadside_market_v1", "橡果集市", "市场", "P0"),
        ("animal_chicken_v1", "鸡", "动物", "P0"),
        ("animal_cow_v1", "牛", "动物", "P0"),
        ("animal_pig_v1", "猪", "后续动物", "P2"),
        ("animal_bear_v1", "熊村民", "居民候选", "P2"),
        ("animal_rabbit_v1", "兔村民", "居民候选", "P2"),
    ]
    add_table(
        document,
        ["资源 ID / 名称", "用途", "规格 / 状态", "负责人 / 优先级"],
        [[asset_id, purpose, f"512 x 512 RGBA；ART-003 READY；NOT_RUNTIME；{category}", f"Codex /root / {priority}"] for asset_id, purpose, category, priority in asset_ids],
        [2.3, 1.3, 2.5, 1.5],
    )
    document.add_paragraph(
        "来源与 SHA、切分矩形、alpha、包围盒、占用率和原点见 `output/art/ART-003-FARM2/split-manifest.json`。"
        "首次工程消费者必须把获选子集复制到新的正式/运行时目录，不得从 candidate 路径加载。"
    )

    document.add_heading("10.2 音乐音效需求", level=2)
    add_table(
        document,
        ["音频 ID / 名称", "触发条件", "播放 / 中断规则", "负责人 / 优先级"],
        [
            ["sfx_crop_swipe", "连续收获/播种逐格成功", "轻短、可叠加但限频；减弱动态不关闭声音", "Audio / P1"],
            ["sfx_collect_product", "动物或机器收取成功", "只在模型提交后播放", "Audio / P1"],
            ["sfx_blocked_storage", "仓满阻塞", "一次提示；快速重试限频", "Audio / P1"],
            ["sfx_queue_add", "配方入队成功", "失败不播放成功音", "Audio / P1"],
            ["sfx_request_complete", "委托事务完成", "奖励和换单事件后播放", "Audio / P1"],
            ["music_town_day", "进入主地图", "循环；设置音量控制", "Audio / P2"],
        ],
        [1.8, 2.1, 2.4, 1.2],
    )

    document.add_heading("10.3 功能打点需求", level=2)
    add_table(
        document,
        ["事件名", "触发时机", "参数 / 数据源", "分析目的"],
        [
            ["crop_planted", "模型提交播种", "crop_id, plot_count, seed_before/after", "验证种子经济与连续操作"],
            ["crop_harvested", "模型提交收获", "crop_id, yield, granary_fill", "验证净增殖与容量压力"],
            ["production_queued", "配方成功入队", "machine_id, recipe_id, queue_depth", "验证队列使用"],
            ["production_blocked", "入队/收取失败", "reason, missing_item, storage_fill", "定位阻塞"],
            ["animal_fed", "喂养提交", "pen_id, feed_cost", "验证饲料闭环"],
            ["request_completed", "委托事务成功", "slot, item_count, coin, renown", "验证目标节奏"],
            ["request_discarded", "玩家丢弃请求", "slot, refresh_seconds", "评估请求质量"],
            ["market_sold", "集市事务成功", "item_id, amount, coin", "验证仓位安全阀"],
            ["camera_pan", "相机拖动结束", "distance, duration", "验证大地图可用性"],
            ["language_changed", "语言保存成功", "from, to", "验证语言入口"],
        ],
        [1.8, 2.0, 2.6, 1.6],
    )

    document.add_heading("11. 关联拓展", level=1)
    add_bullets(
        document,
        [
            "F-004 农场分区与工业 I：依赖 F-003 行为验收；复用双仓、机器、动物和建设合同。",
            "F-005 火车货运：等级解锁后的多箱、多物品、中期订单；不进入当前工程范围。",
            "F-006 城镇服务：访客与服务建筑消费产物；不改变库存权威。",
            "F-007 空运大订单：长周期混合货物批次；需要独立时间和奖励合同。",
            "F-008 土地与分区扩张：清障、地块和建筑网络；沿用 world.csv 与占地元数据。",
            "联网市场、合作社和排行涉及服务器权威；若立项必须使用制作人授权的阿里云档案，不得在本机部署持久服务。",
        ],
    )

    document.add_heading("12. 验收与 QA", level=1)
    document.add_heading("12.1 交付验收", level=2)
    checklist = [
        "☒ 策划：目标、规则、状态、边界、配置和版本一致。",
        "☒ Figma：系统框架、UE、主地图、对象交互、状态和设置有可编辑层及完整预览。",
        "☒ 美术准备：26 项 ART-003 候选具备来源、透明化、尺寸、原点和联系表；仍为 NOT_RUNTIME。",
        "☐ 数据：九张 f003_v2 CSV 已创建并由运行时校验。",
        "☐ 程序：种子经济、双仓、动物、队列、委托和市场全部由权威模型处理。",
        "☐ 地图：720 x 1280 真机画面显示大于一屏世界、拖拽相机和 18+ 可辨对象。",
        "☐ 操作：田地连续操作与相机拖动不冲突。",
        "☐ 边界：缺种、缺料、仓满、队列满、刷新、中断和旧存档无丢失/透支。",
        "☐ 节奏：配置化前 10 分钟模拟始终存在至少一个有意义动作。",
        "☐ 语言：首次 zh-CN；en 可选且重启保持；无未登记硬编码玩家文本。",
        "☐ QA：模型测试、无头启动、真实玩家画面、操作路径和作用域回归通过。",
        "☒ 文档：DOCX/PDF 结构、Figma 登记和 ART-003 证据齐备；逐页视觉 QA 在本回执后归档。",
    ]
    for item in checklist:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("12.2 关键行为测试", level=2)
    add_table(
        document,
        ["ID", "步骤 / 条件", "预期结果", "证据"],
        [
            ["QA-01", "12 个金穗草连续播种 12 格", "库存恰好为 0；第 13 格不透支", "模型测试 + 画面"],
            ["QA-02", "12 格成熟后收获", "获得 36，净增长 24", "模型测试"],
            ["QA-03", "容量不足以容纳一格完整产量", "田格保持成熟；不部分收取", "状态测试 + 画面"],
            ["QA-04", "机器入队失败/成功", "失败不扣料；成功只扣一次", "模型测试"],
            ["QA-05", "成品位未收时跨过多个完成时间", "后续队列不穿透完成", "返场测试"],
            ["QA-06", "货物屋满时收动物/机器", "保持可收取，无丢失", "边界测试 + 画面"],
            ["QA-07", "多物品委托任一不足", "全部不扣；请求不更换", "事务测试"],
            ["QA-08", "丢弃委托后重进", "refresh_at 连续；请求不随机变化", "保存重进测试"],
            ["QA-09", "从空地拖动与从田格划动", "相机/田地手势不互触发", "输入测试 + 录像"],
            ["QA-10", "切后台跨越完成时间", "只推进状态，不自动收取", "中断测试"],
            ["QA-11", "旧 F003-FARM.1 存档迁移", "安全初始种子、无负数、无丢建筑/语言", "迁移测试"],
            ["QA-12", "切换 en 后重启", "英文保持；切回 zh-CN 同样持久", "设置测试 + 画面"],
        ],
        [0.7, 2.8, 2.8, 1.5],
    )

    document.add_heading("12.3 待确认与评审记录", level=2)
    add_table(
        document,
        ["问题 / 决策", "状态", "责任人", "截止日期", "结论及影响版本"],
        [
            ["可编辑 Figma 节点与预览", "已完成", "Codex /root", "2026-07-24", "F003-FARM.2 UI/UE 门禁通过"],
            ["ART-003 资源准备", "已完成 / NOT_RUNTIME", "Codex /root", "2026-07-24", "26 项可供工程回执选择"],
            ["通用模板 DOCX/PDF", "本轮评审", "Codex /root", "2026-07-24", "逐页通过后开启工程只读基线"],
            ["Godot 工程写锁", "待开启", "Codex /root", "后续", "单独回执验证哈希、配置、资源和测试"],
        ],
        [2.0, 1.2, 1.2, 1.1, 2.5],
    )

    document.add_heading("12.4 正式证据索引", level=2)
    add_bullets(
        document,
        [
            "`docs/receipts/F-003-FARM2-DESIGN-001.md`：可编辑 Figma 门禁。",
            "`docs/receipts/ART-003-FARM2-001.md`：来源板、26 项候选、机械与视觉 QA。",
            "`output/figma/F003-FARM.2/city_of_animals_f003_farm2_ue-preview.png`：完整评审预览。",
            "`output/art/ART-003-FARM2/split-manifest.json`：资源哈希、透明度、尺寸、包围盒、占用率和原点。",
            "`PM/feature_progress.xlsx`：当前阶段与后续里程碑。",
        ],
    )

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    document.save(OUTPUT)
    check = Document(OUTPUT)
    all_text = "\n".join(
        [paragraph.text for paragraph in check.paragraphs]
        + [cell.text for table in check.tables for row in table.rows for cell in row.cells]
    )
    forbidden = ["[填写]", "[粘贴", "[YYYY", "[机制名称]", "在此插入"]
    leftovers = [token for token in forbidden if token in all_text]
    if leftovers:
        raise RuntimeError(f"Template placeholders remain: {leftovers}")
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
