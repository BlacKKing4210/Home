from __future__ import annotations

import argparse
import hashlib
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[3]
HELPER_PATH = ROOT / "output/documents/F004-RESIDENT.1/build_f004_resident_doc.py"


def load_helper():
    spec = importlib.util.spec_from_file_location("f004_doc_helper", HELPER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load document helper: {HELPER_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


H = load_helper()


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph(style="Template Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("F-004.2 乳品邻里与居民日常")
    H.set_run_font(run, name="微软雅黑", size=22, bold=True)

    subtitle = doc.add_paragraph(style="Template Note")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "正式功能设计与运行验收包 V1.2 · 720×1280 竖屏 · "
        "RUNTIME_SLICE_APPROVED / SCALE_OUT_APPROVED"
    )
    H.set_run_font(run, name="微软雅黑", size=9.5, color=H.BLUE_DARK)

    memory = doc.add_table(rows=1, cols=1)
    memory.style = "Table Grid"
    H.set_table_geometry(memory, widths=[9360])
    H.set_cell_shading(memory.cell(0, 0), H.GOLD)
    p = memory.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "动物不是按钮或加成，而是玩家看得见、会走路、会生活、"
        "会把小镇运转起来的居民。"
    )
    H.set_run_font(run, name="微软雅黑", size=13.5, bold=True)
    doc.add_paragraph()

    status = doc.add_table(rows=1, cols=1)
    status.style = "Table Grid"
    H.set_table_geometry(status, widths=[9360])
    H.set_cell_shading(status.cell(0, 0), H.GREEN)
    p = status.cell(0, 0).paragraphs[0]
    run = p.add_run(
        "验收结论：F004.2 代表性扩面切片已通过真实 Godot 行为、玩家可见画面、"
        "性能、资产隔离与回归验收。SCALE_OUT_APPROVED 仅批准复用本切片规则，"
        "不代表全量建筑、铁路、空运、海运、Android 真机或发行已完成。"
    )
    H.set_run_font(run, size=9.5, bold=True, color=H.GREEN_DARK)
    doc.add_paragraph()

    H.add_section_table(
        doc,
        "版本控制",
        [
            ["文档编号", "COA-F004-RESIDENT.2", "当前版本", "V1.2 FINAL"],
            ["功能 ID", "F-004.2", "功能状态", "CLOSED"],
            ["唯一制作人", "Codex /root", "运行时状态", "RUNTIME_SLICE_APPROVED"],
            ["创建日期", "2026-07-27", "可复用规则", "SCALE_OUT_APPROVED"],
            ["目标平台", "Mobile / 720×1280 portrait", "默认语言", "zh-CN；设置可切 en"],
            ["可编辑设计", "Penpot / 对象级回读通过", "正式资产", "approved-only / 无候选泄漏"],
            ["行为验证", "91 项 F004.2 PASS", "视觉严重度", "BLOCKER 0 / MATERIAL 0 / POLISH 0"],
        ],
        [1600, 3080, 1600, 3080],
    )
    H.add_section_table(
        doc,
        "正式来源与证据",
        [
            ["类别", "正式来源", "状态"],
            ["制作人决策", "docs/decisions/PD-004-resident-quality-scaleout.md", "Approved"],
            ["功能设计", "docs/features/F-004-resident-dairy-neighborhood-scaleout.md", "V1.2"],
            ["Penpot", "output/penpot/F004-RESIDENT.2/penpot-import-manifest.json", "Readback verified"],
            ["视觉合同", "docs/design/F004-RESIDENT.2-visual-quality-contract.md", "Approved"],
            ["资产清单", "assets/runtime/f004_resident_slice2/runtime-manifest.json", "Approved-only"],
            ["运行验收", "docs/receipts/F004-RESIDENT-SCALEOUT-RUNTIME-ACCEPTANCE-016.md", "Accepted"],
            ["玩家证据", "docs/evidence/F004-RESIDENT.2/README.md", "11 captures"],
        ],
        [1600, 5560, 2200],
    )


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", image_path.stem)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    H.set_run_font(p.add_run(caption), size=8.6, color="666666")


def add_runtime_gallery(doc: Document, images: list[tuple[Path, str]]) -> None:
    for index in range(0, len(images), 2):
        group = images[index : index + 2]
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        H.set_table_geometry(table, widths=[4680, 4680])
        for column in range(2):
            if column >= len(group):
                continue
            image_path, caption = group[column]
            picture_paragraph = table.cell(0, column).paragraphs[0]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = picture_paragraph.add_run().add_picture(str(image_path), width=Cm(7.15))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", image_path.stem)
            caption_paragraph = table.cell(1, column).paragraphs[0]
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            H.set_run_font(caption_paragraph.add_run(caption), size=7.8, color="555555")
        doc.add_paragraph()


def add_executive_summary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("执行摘要与验收边界", style="Heading 1")
    H.add_text_paragraph(
        doc,
        "F004.2 不是静态概念稿，而是一个已经接入正式主入口的可玩代表性扩面切片。"
        "它在不增加高频生产按钮的前提下，把第二住房、第二居民、乳品生产链、"
        "第二辆世界订单车和双居民日常组织成一条可观察的慢节奏闭环。"
    )
    H.add_section_table(
        doc,
        "用户五项要求的落地映射",
        [
            ["用户要求", "F004.2 玩家可见实现", "行为/资产证据"],
            ["统一占地", "1×1 道路生活点、2×2 住房、3×3 牧场、2×2 工坊；重叠、越界和未接路可见", "02、03、04 截图；配置与行为测试"],
            ["世界车辆订单", "车辆到达、等待 0/1、装载 1/1、鸣笛离场并一次结算 120 金币", "08、09 截图；settlement 幂等断言"],
            ["正式质感", "道路、田地/牧场、建筑、车辆、居民均使用原创 approved 资产；主页面无临时字母块", "运行清单与 candidate leak=false"],
            ["居民自主作业", "熊沿道路完成照料、挤奶、搬运、加工和装车；玩家不逐批点击生产", "05–09 截图；任务状态机断言"],
            ["慢节奏居民生活", "建房、邀请、入住、长期派遣；兔子与熊工作后回家或停留生活点", "04、10、11 截图；存档/中断恢复"],
        ],
        [1900, 4540, 2920],
    )
    H.add_section_table(
        doc,
        "真实验证摘要",
        [
            ["项目", "结果"],
            ["F004.2 行为断言", "91 PASS / F004_RESIDENT2_TESTS_PASSED"],
            ["回归", "F004.1、F003 FARM.2、镇区模型全部 PASS"],
            ["主入口", "project.godot → scenes/town_main.tscn → f004_resident2_view.gd"],
            ["真实画面", "11 张 720×1280 Windows OpenGL 玩家视角"],
            ["性能", "180 帧实测约 167 FPS；捕获时约 101 FPS；静态内存约 50.7 MB"],
            ["资产隔离", "12 个 approved/shared 运行资源；candidate/source/qa 泄漏=false"],
            ["可访问性", "zh-CN 默认、en 可切换并持久化、reduced-motion"],
            ["严重度", "BLOCKER 0 / MATERIAL 0 / POLISH 0"],
        ],
        [2500, 6860],
    )
    H.add_text_paragraph(
        doc,
        "明确边界：本验收不包含全量历史建筑重制、铁路、空运、海运、复杂社交、"
        "Android 真机、发行包或商业化。后续扩面必须创建独立 feature identity、"
        "RAG task receipt、只读回执与精确写集，不能把本切片的通过直接外推为整款游戏完成。",
        style="Template Note",
        color=H.BLUE_DARK,
        size=9.5,
    )


def build(template: Path, output: Path) -> None:
    doc = Document(str(template))
    H.clear_document_body(doc)
    H.style_document(doc)
    H.add_footer(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    add_cover(doc)
    doc.add_page_break()
    doc.add_paragraph("目录", style="Heading 1")
    H.add_field_toc(doc.add_paragraph())
    H.add_text_paragraph(
        doc,
        "目录在 Word 打开或 PDF 导出时更新，覆盖 1–3 级标题。",
        style="Template Note",
        color=H.BLUE_DARK,
        size=9.2,
    )

    add_executive_summary(doc)

    doc.add_page_break()
    doc.add_paragraph("Penpot 可编辑设计与状态流程", style="Heading 1")
    H.add_text_paragraph(
        doc,
        "下列画板和流程已经写入认证 Penpot 文件并在保存、重新打开后完成对象级回读。"
        "Penpot 是本功能 UI/UE 的可编辑变更源；本地 SVG 是导入与备份源，PNG/PDF 只是审阅预览。"
    )
    add_figure(
        doc,
        ROOT / "output/penpot/F004-RESIDENT.2/previews/F004-RESIDENT.2-screen-board.png",
        "图 1：六个 720×1280 Penpot 画面状态，覆盖放置、非法占地、邀请、派遣、工作车辆和双居民日常。",
        16.4,
    )
    add_figure(
        doc,
        ROOT / "output/penpot/F004-RESIDENT.2/previews/F004-RESIDENT.2-flow-board.png",
        "图 2：放置合法性、熊居民乳品状态机、第二车辆订单和双居民日程流程。",
        16.4,
    )
    H.parse_markdown(
        doc,
        (ROOT / "output/penpot/F004-RESIDENT.2/README.md").read_text(encoding="utf-8"),
    )

    doc.add_page_break()
    doc.add_paragraph("真实 720×1280 运行证据", style="Heading 1")
    H.add_text_paragraph(
        doc,
        "以下图片来自 Godot 正式渲染路径和主入口状态，不是 Penpot 或静态合成图。"
        "每张截图与行为断言、存档状态和资源清单共同构成验收证据。"
    )
    runtime_dir = ROOT / "output/runtime/F004-RESIDENT.2"
    gallery = [
        (runtime_dir / "01-default-neighborhood.png", "01 默认邻里：世界车辆到达，原住房、兔子、道路和装卸院可见。"),
        (runtime_dir / "02-invalid-overlap.png", "02 非法重叠：红色占地、冲突图形和禁用确认。"),
        (runtime_dir / "03-valid-road-footprint.png", "03 合法 1×1 道路生活点：绿色完整占地与连接。"),
        (runtime_dir / "04-neighborhood-invited.png", "04 新邻里：2×2 住房、3×3 牧场、2×2 工坊、熊与奶牛。"),
        (runtime_dir / "05-road-interrupted.png", "05 道路中断：居民在安全点等待，可定位缺口。"),
        (runtime_dir / "06-pasture-work.png", "06 牧场作业：熊照料奶牛，五段图形轨迹推进。"),
        (runtime_dir / "07-creamery-work.png", "07 工坊作业：熊在世界内加工乳品。"),
        (runtime_dir / "08-loading-blocked.png", "08 装卸阻塞：货物保留，不吞物品、不重复产出。"),
        (runtime_dir / "09-vehicle-departing.png", "09 车辆离场：1/1 装载并一次结算 120 金币。"),
        (runtime_dir / "10-two-residents-life.png", "10 双居民生活：熊和兔子回到生活点，慢节奏闭环完成。"),
        (runtime_dir / "11-settings-en-reduced.png", "11 英文与减少动态：设置持久化，无成功粒子泄漏。"),
    ]
    add_runtime_gallery(doc, gallery)

    sections = [
        ("正式功能设计", "docs/features/F-004-resident-dairy-neighborhood-scaleout.md"),
        ("主页面 UI/UX 优先级", "docs/uiux/F004-RESIDENT.2-ui-priority.md"),
        ("视觉质量合同", "docs/design/F004-RESIDENT.2-visual-quality-contract.md"),
        ("正式资产合同", "docs/art/F004-RESIDENT.2-asset-contract.md"),
        ("玩家可见证据索引", "docs/evidence/F004-RESIDENT.2/README.md"),
        ("运行时视觉对比清单", "docs/evidence/F004-RESIDENT.2/runtime-parity-checklist.md"),
        ("运行验收回执", "docs/receipts/F004-RESIDENT-SCALEOUT-RUNTIME-ACCEPTANCE-016.md"),
        ("制作人扩面决策", "docs/decisions/PD-004-resident-quality-scaleout.md"),
    ]
    for title, relative_path in sections:
        heading = doc.add_paragraph(f"附录：{title}", style="Heading 1")
        heading.paragraph_format.page_break_before = True
        H.parse_markdown(doc, (ROOT / relative_path).read_text(encoding="utf-8"))

    H.ensure_update_fields(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"built={output}")
    print(f"sha256={hashlib.sha256(output.read_bytes()).hexdigest()}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.template, args.output)


if __name__ == "__main__":
    main()
