from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


INK = "#243238"
PAPER = "#FFF8E7"
BLUE = "#D8E4F2"
BLUE_DARK = "#54789A"
GREEN = "#BFD9A8"
GREEN_DARK = "#5F8A56"
GOLD = "#F2D27C"
CORAL = "#E89B82"
TEAL = "#74B7AE"
GRAY = "#E9ECEF"
RED = "#C85A54"
WHITE = "#FFFFFF"


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font_path() -> str:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    raise FileNotFoundError("No Chinese font found")


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = r"C:\Windows\Fonts\msyhbd.ttc" if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists() else FONT_PATH
    return ImageFont.truetype(path, size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_label(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    *,
    fill: str,
    outline: str = INK,
    radius: int = 20,
    font_size: int = 30,
    bold: bool = False,
    text_color: str = INK,
    padding: int = 18,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    font = pil_font(font_size, bold)
    max_width = max(20, box[2] - box[0] - padding * 2)
    lines = wrap_text(draw, text, font, max_width)
    line_height = font_size + 10
    total_height = len(lines) * line_height
    y = box[1] + (box[3] - box[1] - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=text_color)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = INK, width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    p1 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p2 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([end, p1, p2], fill=color)


def add_diagram_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str, width: int) -> None:
    draw.text((56, 42), title, font=pil_font(48, True), fill=hex_rgb(INK))
    draw.text((58, 108), subtitle, font=pil_font(24), fill=hex_rgb(BLUE_DARK))
    draw.line([(56, 150), (width - 56, 150)], fill=hex_rgb(INK), width=3)


def add_draft_mark(draw: ImageDraw.ImageDraw, width: int, height: int) -> None:
    text = "本地低保真评审草图｜Figma 可编辑门禁未完成"
    font = pil_font(22)
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((width - 48 - (bbox[2] - bbox[0]), height - 44), text, font=font, fill=hex_rgb(RED))


def build_system_framework(path: Path) -> None:
    w, h = 1800, 1000
    img = Image.new("RGB", (w, h), hex_rgb(PAPER))
    d = ImageDraw.Draw(img)
    add_diagram_header(d, "F-004 系统框架", "玩家从大地图识别分区，通过营建板打开固定站点并形成一级工业网络", w)
    boxes = {
        "地图": (70, 220, 360, 360),
        "营建板": (450, 220, 740, 360),
        "目标": (830, 220, 1120, 360),
        "站点": (1210, 220, 1500, 360),
        "来源": (170, 500, 460, 640),
        "动物": (565, 500, 855, 640),
        "机器": (960, 500, 1250, 640),
        "库存": (1355, 500, 1645, 640),
        "订单": (755, 760, 1045, 900),
    }
    fills = [GREEN, GOLD, BLUE, CORAL, GREEN, TEAL, CORAL, BLUE, GOLD]
    labels = [
        "1800×1700 大地图\n三分区与16个功能入口",
        "分区营建板\n一次只突出一个下一步",
        "一次性目标\n完成后原子发奖",
        "固定功能位\n锁定→可建→建造→建成",
        "田地与果园\n手动收取",
        "猪舍与兔苑\n喂养→计时→收取",
        "合织/压榨/保鲜\n队列与成品台",
        "共享库存\n容量与待领位",
        "既有委托/市场\n后续货运接口",
    ]
    for (key, box), fill, label in zip(boxes.items(), fills, labels):
        draw_label(d, box, label, fill=fill, font_size=27, bold=key in {"地图", "营建板", "目标", "站点"})
    for a, b in [("地图", "营建板"), ("营建板", "目标"), ("目标", "站点")]:
        aa, bb = boxes[a], boxes[b]
        draw_arrow(d, (aa[2], (aa[1] + aa[3]) // 2), (bb[0], (bb[1] + bb[3]) // 2), BLUE_DARK)
    for a in ["来源", "动物", "机器"]:
        aa, bb = boxes["站点"], boxes[a]
        draw_arrow(d, ((aa[0] + aa[2]) // 2, aa[3]), ((bb[0] + bb[2]) // 2, bb[1]), GREEN_DARK)
    for a in ["来源", "动物", "机器"]:
        aa, bb = boxes[a], boxes["库存"]
        draw_arrow(d, (aa[2], (aa[1] + aa[3]) // 2), (bb[0], (bb[1] + bb[3]) // 2), TEAL)
    draw_arrow(d, ((boxes["库存"][0] + boxes["库存"][2]) // 2, boxes["库存"][3]), ((boxes["订单"][0] + boxes["订单"][2]) // 2, boxes["订单"][1]), CORAL)
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def build_ue_flow(path: Path) -> None:
    w, h = 2000, 1040
    img = Image.new("RGB", (w, h), hex_rgb(PAPER))
    d = ImageDraw.Draw(img)
    add_diagram_header(d, "玩家 UE 主流程", "主路径保持对象优先；错误、容量不足与中断均回到可恢复状态", w)
    nodes = [
        ("进入大地图", GREEN),
        ("点击营建板", GOLD),
        ("选择分区", BLUE),
        ("查看唯一下一步", BLUE),
        ("定位目标或工地", GREEN),
        ("达成目标/开始建造", CORAL),
        ("进入生产上下文", TEAL),
        ("手动收取并回地图", GOLD),
    ]
    x0, y0 = 50, 230
    bw, bh, gap = 205, 150, 30
    boxes = []
    for idx, (label, fill) in enumerate(nodes):
        x = x0 + idx * (bw + gap)
        box = (x, y0, x + bw, y0 + bh)
        boxes.append(box)
        draw_label(d, box, f"{idx + 1}. {label}", fill=fill, font_size=25, bold=idx in {0, 3, 7})
        if idx:
            draw_arrow(d, (boxes[idx - 1][2], y0 + bh // 2), (box[0], y0 + bh // 2), BLUE_DARK, 5)
    branches = [
        ("资源不足", "不扣资源；突出缺口；返回站点", (170, 550, 570, 710), RED),
        ("仓储已满", "成品留在成品台；允许先清仓再收取", (610, 550, 1010, 710), GOLD),
        ("中断/重进", "推进已开始计时；不自动收取、不重复发奖", (1050, 550, 1450, 710), TEAL),
        ("配置错误", "禁用入口；记录诊断；保持存档可读", (1490, 550, 1890, 710), GRAY),
    ]
    for title, text, box, fill in branches:
        draw_label(d, box, f"{title}\n{text}", fill=fill, font_size=24, bold=True)
    draw_arrow(d, ((boxes[5][0] + boxes[5][2]) // 2, boxes[5][3]), (370, 550), RED, 4)
    draw_arrow(d, ((boxes[7][0] + boxes[7][2]) // 2, boxes[7][3]), (810, 550), GOLD, 4)
    draw_arrow(d, ((boxes[6][0] + boxes[6][2]) // 2, boxes[6][3]), (1250, 550), TEAL, 4)
    d.text((70, 820), "返回原则：关闭面板或失败后回到原地图焦点；玩家已选择的分区保持不变。", font=pil_font(28, True), fill=hex_rgb(INK))
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def building(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, fill: str, roof: str, label: str, bubble: str | None = None) -> None:
    draw.rounded_rectangle((x, y + 25, x + w, y + h), radius=16, fill=hex_rgb(fill), outline=hex_rgb(INK), width=4)
    draw.polygon([(x - 8, y + 40), (x + w // 2, y), (x + w + 8, y + 40)], fill=hex_rgb(roof), outline=hex_rgb(INK))
    f = pil_font(19, True)
    bbox = draw.textbbox((0, 0), label, font=f)
    draw.text((x + (w - (bbox[2] - bbox[0])) / 2, y + h - 34), label, font=f, fill=hex_rgb(INK))
    if bubble:
        draw.ellipse((x + w - 20, y - 30, x + w + 32, y + 22), fill=hex_rgb(WHITE), outline=hex_rgb(INK), width=3)
        fb = pil_font(20, True)
        bb = draw.textbbox((0, 0), bubble, font=fb)
        draw.text((x + w + 6 - (bb[2] - bb[0]) / 2, y - 19), bubble, font=fb, fill=hex_rgb(RED))


def build_main_map(path: Path) -> None:
    w, h = 720, 1280
    img = Image.new("RGB", (w, h), hex_rgb("#CFE7B4"))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((20, 20, 700, 105), radius=24, fill=hex_rgb(PAPER), outline=hex_rgb(INK), width=4)
    d.text((44, 43), "Lv 9", font=pil_font(28, True), fill=hex_rgb(INK))
    d.text((210, 43), "金币 460", font=pil_font(26, True), fill=hex_rgb(INK))
    d.text((440, 43), "粮仓 31/60", font=pil_font(24, True), fill=hex_rgb(INK))
    d.rounded_rectangle((18, 125, 702, 1125), radius=28, fill=hex_rgb("#EAF2D4"), outline=hex_rgb(INK), width=4)
    d.polygon([(18, 125), (390, 125), (255, 700), (18, 760)], fill=hex_rgb("#D7EDB7"))
    d.polygon([(390, 125), (702, 125), (702, 720), (255, 700)], fill=hex_rgb("#D2E8DF"))
    d.polygon([(18, 760), (255, 700), (702, 720), (702, 1125), (18, 1125)], fill=hex_rgb("#F5E1C3"))
    d.text((42, 150), "丰收花园", font=pil_font(23, True), fill=hex_rgb(GREEN_DARK))
    d.text((500, 150), "草坡牧场", font=pil_font(23, True), fill=hex_rgb(BLUE_DARK))
    d.text((430, 890), "匠作长街", font=pil_font(23, True), fill=hex_rgb("#9B5D45"))
    for row in range(2):
        for col in range(3):
            x = 48 + col * 92
            y = 230 + row * 84
            d.rectangle((x, y, x + 72, y + 60), fill=hex_rgb(GOLD if (row + col) % 2 == 0 else GREEN), outline=hex_rgb(INK), width=2)
            d.line((x + 10, y + 30, x + 62, y + 30), fill=hex_rgb(GREEN_DARK), width=3)
    building(d, 95, 440, 120, 120, "#C89F74", GOLD, "粮仓", "!")
    building(d, 285, 300, 125, 128, "#D6B678", CORAL, "果园", "✓")
    building(d, 480, 300, 130, 128, "#B8D6CE", TEAL, "兔苑", "●")
    building(d, 445, 510, 145, 132, "#DAB38E", BLUE_DARK, "猪舍", None)
    building(d, 90, 745, 145, 132, "#E1C49C", CORAL, "混粮坊", "…")
    building(d, 285, 760, 145, 132, "#D4B1A2", RED, "合织坊", "!")
    building(d, 500, 760, 145, 132, "#C8D5A0", GOLD, "压榨坊", "✓")
    building(d, 250, 950, 180, 130, "#E6D6A6", BLUE_DARK, "分区营建板", "→")
    d.rounded_rectangle((18, 1145, 702, 1260), radius=28, fill=hex_rgb(PAPER), outline=hex_rgb(INK), width=4)
    labels = ["地图", "仓库", "订单", "商店"]
    for i, label in enumerate(labels):
        x = 48 + i * 165
        d.rounded_rectangle((x, 1170, x + 120, 1238), radius=18, fill=hex_rgb(GOLD if i == 0 else GRAY), outline=hex_rgb(INK), width=2)
        d.text((x + 30, 1188), label, font=pil_font(22, True), fill=hex_rgb(INK))
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def build_district_board(path: Path) -> None:
    w, h = 720, 1280
    img = Image.new("RGB", (w, h), hex_rgb(PAPER))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((18, 20, 702, 1260), radius=30, fill=hex_rgb("#FFFDF6"), outline=hex_rgb(INK), width=4)
    d.text((44, 48), "分区营建板", font=pil_font(38, True), fill=hex_rgb(INK))
    d.text((596, 52), "关闭", font=pil_font(22, True), fill=hex_rgb(BLUE_DARK))
    cards = [
        ("丰收花园", "3/4", GREEN),
        ("草坡牧场", "2/4", TEAL),
        ("匠作长街", "1/4", CORAL),
    ]
    for i, (name, progress, fill) in enumerate(cards):
        x = 38 + i * 220
        draw_label(d, (x, 130, x + 202, 260), f"{name}\n{progress}", fill=fill, font_size=24, bold=i == 0)
    d.text((42, 310), "当前下一步", font=pil_font(28, True), fill=hex_rgb(INK))
    d.rounded_rectangle((38, 355, 682, 575), radius=24, fill=hex_rgb(BLUE), outline=hex_rgb(INK), width=4)
    d.ellipse((70, 395, 188, 513), fill=hex_rgb(GOLD), outline=hex_rgb(INK), width=3)
    d.text((105, 427), "麦", font=pil_font(42, True), fill=hex_rgb(GREEN_DARK))
    d.text((220, 390), "累计收取 4 份金穗草", font=pil_font(29, True), fill=hex_rgb(INK))
    d.rounded_rectangle((220, 452, 630, 487), radius=16, fill=hex_rgb(WHITE), outline=hex_rgb(INK), width=2)
    d.rounded_rectangle((220, 452, 520, 487), radius=16, fill=hex_rgb(GREEN), outline=hex_rgb(INK), width=2)
    d.text((220, 510), "进度 3/4  ·  奖励：木材与金币", font=pil_font(22), fill=hex_rgb(INK))
    d.text((42, 620), "固定功能位", font=pil_font(28, True), fill=hex_rgb(INK))
    rows = [
        ("琥珀果园", "可建造", GOLD),
        ("绒鬃猪舍", "目标锁定", GRAY),
        ("软绒兔苑", "建造中 00:28", TEAL),
        ("合织工坊", "前置不足", CORAL),
    ]
    for i, (name, state, fill) in enumerate(rows):
        y = 670 + i * 105
        d.rounded_rectangle((40, y, 680, y + 82), radius=18, fill=hex_rgb(fill), outline=hex_rgb(INK), width=3)
        d.text((68, y + 23), name, font=pil_font(25, True), fill=hex_rgb(INK))
        d.text((410, y + 24), state, font=pil_font(22), fill=hex_rgb(INK))
    d.rounded_rectangle((105, 1125, 615, 1215), radius=28, fill=hex_rgb(CORAL), outline=hex_rgb(INK), width=4)
    d.text((240, 1147), "前往查看", font=pil_font(32, True), fill=hex_rgb(INK))
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def build_production_states(path: Path) -> None:
    w, h = 1800, 1100
    img = Image.new("RGB", (w, h), hex_rgb(PAPER))
    d = ImageDraw.Draw(img)
    add_diagram_header(d, "建造与生产上下文", "建筑点击后打开底栏；成品台、队列、配方与主动作分区明确", w)
    panels = [
        ("果园", "可收取", GREEN, ["成品台：琥珀果", "下一轮：70 秒", "主动作：收取"]),
        ("动物栏", "货物屋已满", TEAL, ["成品台：轻绒纤维", "状态：可收取", "主动作：先清仓"]),
        ("工坊", "队列 2/3", CORAL, ["成品台：1/3", "队列：进行中/等待", "主动作：加入队列"]),
    ]
    for i, (name, state, fill, details) in enumerate(panels):
        x = 55 + i * 575
        d.rounded_rectangle((x, 210, x + 520, 930), radius=26, fill=hex_rgb("#FFFDF8"), outline=hex_rgb(INK), width=4)
        d.rounded_rectangle((x + 25, 245, x + 495, 335), radius=20, fill=hex_rgb(fill), outline=hex_rgb(INK), width=3)
        d.text((x + 48, 268), f"{name}｜{state}", font=pil_font(29, True), fill=hex_rgb(INK))
        d.text((x + 36, 380), "成品台", font=pil_font(24, True), fill=hex_rgb(INK))
        for j in range(3):
            d.rounded_rectangle((x + 38 + j * 145, 425, x + 155 + j * 145, 535), radius=18, fill=hex_rgb(GOLD if j == 0 else GRAY), outline=hex_rgb(INK), width=3)
            if j == 0:
                d.text((x + 78, 458), "成品", font=pil_font(22, True), fill=hex_rgb(INK))
        d.text((x + 36, 570), "生产队列", font=pil_font(24, True), fill=hex_rgb(INK))
        for j in range(3):
            d.rounded_rectangle((x + 38 + j * 145, 615, x + 155 + j * 145, 700), radius=16, fill=hex_rgb(TEAL if j < 2 else WHITE), outline=hex_rgb(INK), width=3)
        d.text((x + 38, 732), "配方图标  ×  原料数量", font=pil_font(23), fill=hex_rgb(INK))
        d.rounded_rectangle((x + 88, 805, x + 432, 880), radius=24, fill=hex_rgb(fill), outline=hex_rgb(INK), width=3)
        d.text((x + 154, 826), details[-1], font=pil_font(25, True), fill=hex_rgb(INK))
    d.text((60, 975), "气泡优先级：仓储阻挡 > 可收取 > 可建造 > 建造中 > 生产中 > 空闲；每建筑常态最多 1 个。", font=pil_font(27, True), fill=hex_rgb(RED))
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def build_edge_states(path: Path) -> None:
    w, h = 1800, 1100
    img = Image.new("RGB", (w, h), hex_rgb(PAPER))
    d = ImageDraw.Draw(img)
    add_diagram_header(d, "全状态与边界", "每个状态给出玩家可见反馈、允许动作与恢复路径", w)
    states = [
        ("加载中", "骨架位；禁止重复提交", BLUE),
        ("空状态", "保留地图；隐藏无效 CTA", GRAY),
        ("禁用", "显示缺口；点击只解释", GOLD),
        ("失败", "不扣资源；提供重试", CORAL),
        ("成功", "原位变化；短反馈后回地图", GREEN),
        ("中断", "重进推进计时；不自动收取", TEAL),
        ("语言切换", "立即刷新 zh-CN/en；持久化", BLUE),
        ("减弱动态", "镜头瞬移；取消循环装饰", GRAY),
    ]
    for i, (title, text, fill) in enumerate(states):
        col, row = i % 4, i // 4
        x, y = 55 + col * 430, 220 + row * 350
        draw_label(d, (x, y, x + 390, y + 270), f"{title}\n{text}", fill=fill, font_size=27, bold=True)
    d.text((60, 930), "统一回退：保持当前分区选择与原地图焦点；任何失败都不得静默丢失产品、奖励或已开始计时。", font=pil_font(28, True), fill=hex_rgb(INK))
    add_draft_mark(d, w, h)
    img.save(path, quality=95)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, total_twips: int, widths: list[int] | None = None) -> None:
    cols = len(table.columns)
    if widths is None:
        widths = [total_twips // cols] * cols
        widths[-1] += total_twips - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, name: str = "宋体", size: float = 10.5, bold: bool | None = None, color: str = INK) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(*hex_rgb(color))
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1（\2）", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_field_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "打开 Word 后更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])


def add_footer_page_fields(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for child in list(p._p):
        p._p.remove(child)
    run = p.add_run()
    set_run_font(run, size=9, color="#555555")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])
    r2 = p.add_run(" / ")
    set_run_font(r2, size=9, color="#555555")
    r3 = p.add_run()
    set_run_font(r3, size=9, color="#555555")
    b2 = OxmlElement("w:fldChar")
    b2.set(qn("w:fldCharType"), "begin")
    i2 = OxmlElement("w:instrText")
    i2.set(qn("xml:space"), "preserve")
    i2.text = " NUMPAGES "
    s2 = OxmlElement("w:fldChar")
    s2.set(qn("w:fldCharType"), "separate")
    t2 = OxmlElement("w:t")
    t2.text = "1"
    e2 = OxmlElement("w:fldChar")
    e2.set(qn("w:fldCharType"), "end")
    r3._r.extend([b2, i2, s2, t2, e2])


def style_doc(doc: Document) -> None:
    for style_name, size, bold, color in [
        ("Normal", 10.5, False, INK),
        ("Heading 1", 22, True, INK),
        ("Heading 2", 16, True, INK),
        ("Heading 3", 12, True, INK),
        ("List Bullet", 10.5, False, INK),
        ("List Number", 10.5, False, INK),
    ]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(*hex_rgb(color))
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 1"].paragraph_format.keep_with_next = True
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 2"].paragraph_format.keep_with_next = True
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 3"].paragraph_format.keep_with_next = True


def add_title_table(doc: Document, title: str, rows: list[list[str]], widths: list[int]) -> None:
    t0 = doc.add_table(rows=1, cols=1)
    t0.style = "Table Grid"
    set_table_geometry(t0, 9360, [9360])
    set_cell_shading(t0.cell(0, 0), BLUE)
    p = t0.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, name="微软雅黑", size=12, bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, 9360, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=9.2, bold=ridx == 0)
    doc.add_paragraph()


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    if cols >= 3:
        widths[0] = int(9360 * 0.18)
        remaining = 9360 - widths[0]
        widths[1:] = [remaining // (cols - 1)] * (cols - 1)
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, 9360, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            value = clean_inline(row[cidx]) if cidx < len(row) else ""
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=8.3, bold=ridx == 0)
    doc.add_paragraph()


def add_review_image(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, color="#666666")


def create_restarting_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num_id = style.find(qn("w:pPr"))
    if style_num_id is not None:
        style_num_id = style_num_id.find(qn("w:numPr"))
    if style_num_id is not None:
        style_num_id = style_num_id.find(qn("w:numId"))
    base_num_id = int(style_num_id.get(qn("w:val"))) if style_num_id is not None else 5
    base_num = numbering.find(f".//{qn('w:num')}[@{qn('w:numId')}='{base_num_id}']")
    if base_num is None:
        base_num = numbering.find(qn("w:num"))
    abstract = base_num.find(qn("w:abstractNumId"))
    abstract_id = int(abstract.get(qn("w:val")))
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def parse_markdown(doc: Document, markdown: str, image_map: dict[str, tuple[Path, float, str]]) -> None:
    lines = markdown.splitlines()
    start = 0
    for idx, line in enumerate(lines):
        if line.startswith("## 1. "):
            start = idx
            break
    lines = lines[start:]
    i = 0
    in_code = False
    code_kind = ""
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []
    active_num_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = clean_inline(" ".join(x.strip() for x in paragraph_buffer if x.strip()))
        if text:
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            set_run_font(r)
        paragraph_buffer = []

    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if line.startswith("```"):
            flush_paragraph()
            active_num_id = None
            if not in_code:
                in_code = True
                code_kind = line[3:].strip()
                code_lines = []
            else:
                if code_kind == "mermaid":
                    p = doc.add_paragraph(style="Template Note" if "Template Note" in doc.styles else "Normal")
                    r = p.add_run("本段 Mermaid 仅作为逻辑草稿；对应低保真评审图已插入本章，正式可编辑源仍受 Figma 门禁约束。")
                    set_run_font(r, size=9.2, color="#666666")
                elif code_lines:
                    p = doc.add_paragraph(style="Normal")
                    r = p.add_run("\n".join(code_lines))
                    set_run_font(r, name="Consolas", size=8.5, color="#2B5D36")
                in_code = False
                code_kind = ""
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line:
            flush_paragraph()
            active_num_id = None
            i += 1
            continue
        if line == "---":
            flush_paragraph()
            active_num_id = None
            i += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            active_num_id = None
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows: list[list[str]] = []
            for tline in table_lines:
                parts = [p.strip() for p in tline.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
                    continue
                rows.append(parts)
            add_markdown_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            active_num_id = None
            hashes, title = heading_match.groups()
            title = clean_inline(title)
            level = {2: 1, 3: 2, 4: 3}[len(hashes)]
            p = doc.add_paragraph(title, style=f"Heading {level}")
            if title in image_map:
                image_path, width_cm, caption = image_map[title]
                add_review_image(doc, image_path, width_cm, caption)
            i += 1
            continue
        num_match = re.match(r"^\d+\.\s+(.+)$", line)
        if num_match:
            flush_paragraph()
            if active_num_id is None:
                active_num_id = create_restarting_numbering(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_num_id)
            r = p.add_run(clean_inline(num_match.group(1)))
            set_run_font(r)
            i += 1
            continue
        if line.startswith("- "):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean_inline(line[2:]))
            set_run_font(r)
            i += 1
            continue
        if line.startswith(">"):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph(style="Template Note" if "Template Note" in doc.styles else "Normal")
            r = p.add_run(clean_inline(line.lstrip(">").strip()))
            set_run_font(r, size=9.5, color="#666666")
            i += 1
            continue
        active_num_id = None
        paragraph_buffer.append(line)
        i += 1
    flush_paragraph()


def build_document(reference: Path, markdown_path: Path, output: Path, assets_dir: Path) -> None:
    assets_dir.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "framework": assets_dir / "f004-system-framework-draft.png",
        "flow": assets_dir / "f004-player-ue-flow-draft.png",
        "map": assets_dir / "f004-main-map-720x1280-draft.png",
        "board": assets_dir / "f004-district-board-720x1280-draft.png",
        "production": assets_dir / "f004-production-states-draft.png",
        "edges": assets_dir / "f004-edge-states-draft.png",
    }
    build_system_framework(diagrams["framework"])
    build_ue_flow(diagrams["flow"])
    build_main_map(diagrams["map"])
    build_district_board(diagrams["board"])
    build_production_states(diagrams["production"])
    build_edge_states(diagrams["edges"])

    doc = Document(str(reference))
    clear_document_body(doc)
    style_doc(doc)
    add_footer_page_fields(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.style = "Template Title" if "Template Title" in doc.styles else "Heading 1"
    r = title.add_run("F-004 动物小镇农场分区与一级工业")
    set_run_font(r, name="微软雅黑", size=26, bold=True)
    note = doc.add_paragraph()
    note.style = "Template Note" if "Template Note" in doc.styles else "Normal"
    r = note.add_run("研究驱动设计 V1.0｜配置与文档已验证｜运行时未授权｜Figma UE 附件阻塞")
    set_run_font(r, size=10, color="#666666")

    add_title_table(
        doc,
        "版本控制",
        [
            ["文档编号", "CityOfAnimals-F004-V1.0", "当前版本", "V1.0"],
            ["功能 ID", "F004-DISTRICT.1", "文档状态", "配置/文档通过 / Figma 阻塞"],
            ["主策划", "Codex /root", "制作策划", "Codex /root"],
            ["制作程序", "未登记 / 未授权", "制作美术", "未登记 / 未授权"],
            ["创建日期", "2026-07-24", "计划完成", "Figma 恢复后"],
            ["Figma 源", "https://www.figma.com/design/uU2Oek5RqFb19CPoGl48lC/Untitled", "配置版本", "F004 CSV V1.0"],
            ["备注", "720×1280 竖屏；zh-CN 默认，en 可选", "评审结论", "配置/DOCX/PDF 通过；UI/UE 阻塞"],
        ],
        [1600, 3280, 1600, 2880],
    )
    add_title_table(
        doc,
        "版本历史",
        [
            ["版本", "编写人", "审核人", "批准人", "日期", "更新内容"],
            ["V0.8", "Codex /root", "制作人自审", "Pending", "2026-07-24", "三分区、固定功能位、一级工业网络"],
            ["V0.9", "Codex /root", "待里程碑评审", "Pending", "2026-07-24", "完整 UE、状态机、8 张 CSV、保存契约"],
            ["V1.0", "Codex /root", "制作人自审", "HOLD / Figma", "2026-07-25", "研究修订；8 表验证与 27 页 DOCX/PDF QA 通过"],
        ],
        [900, 1350, 1350, 1150, 1350, 3260],
    )
    add_title_table(
        doc,
        "开发计划",
        [
            ["责任方", "负责人", "开始日期", "预计天数", "完成条件"],
            ["策划/配置", "Codex /root", "2026-07-24", "已完成当前批次", "规则、8 张 CSV 和交叉引用审计通过"],
            ["UI/UE", "Codex /root", "2026-07-25", "连接恢复后 1 天", "6 个可编辑 Figma 节点与导出一致"],
            ["程序", "未登记", "未开始", "Pending", "新工程只读回执、精确写集和行为验收"],
            ["QA", "未登记", "未开始", "Pending", "真实 720×1280 玩家可见证据与回归通过"],
        ],
        [1350, 1800, 1500, 1500, 3210],
    )

    doc.add_page_break()
    toc_title = doc.add_paragraph("目录", style="Heading 1")
    toc = doc.add_paragraph()
    add_field_toc(toc)
    p = doc.add_paragraph()
    r = p.add_run("目录字段已设置为 1-3 级标题；PDF 导出前由 Word 更新。")
    set_run_font(r, size=9.5, color="#666666")
    doc.add_page_break()

    image_map = {
        "4.1 模块关系": (diagrams["framework"], 17.0, "图 4-1　本地系统框架评审草图；正式可编辑 Figma 节点 00_System_Framework 仍为 BLOCKED。"),
        "5. 玩家 UE 流程": (diagrams["flow"], 17.0, "图 5-1　本地玩家 UE 主流程评审草图；正式可编辑 Figma 节点 01_Player_UE_Flow 仍为 BLOCKED。"),
        "9.3 主地图与对象状态气泡（720 x 1280）": (diagrams["map"], 9.0, "图 9-1　720×1280 主地图低保真评审草图：对象密度、分区与单气泡优先级。"),
        "9.4 分区营建板布局（720 x 1280）": (diagrams["board"], 9.0, "图 9-2　720×1280 分区营建板低保真评审草图：三分区快照与单主 CTA。"),
        "9.5 生产建筑上下文底栏（720 x 1280）": (diagrams["production"], 17.0, "图 9-3　生产上下文低保真评审草图：成品台、队列、配方与主动作。"),
        "9.7 全状态规范": (diagrams["edges"], 17.0, "图 9-4　状态与边界低保真评审草图：加载、空、禁用、失败、成功与中断。"),
    }
    parse_markdown(doc, markdown_path.read_text(encoding="utf-8"), image_map)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--assets", type=Path, required=True)
    args = parser.parse_args()
    build_document(args.reference, args.markdown, args.output, args.assets)
    digest = hashlib.sha256(args.output.read_bytes()).hexdigest()
    print(f"built={args.output}")
    print(f"sha256={digest}")


if __name__ == "__main__":
    main()
